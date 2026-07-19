"""Generate project-synopsis.docx for Sajilo Inventory.

Usage: python generate_synopsis_docx.py
Output: ../project-synopsis.docx
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "project-synopsis.docx")

doc = Document()

# ── Global style defaults ──────────────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Times New Roman"
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.paragraph_format.space_before = Pt(12)
    hs.paragraph_format.space_after = Pt(6)
    if level == 1:
        hs.font.size = Pt(16)
        hs.font.bold = True
    elif level == 2:
        hs.font.size = Pt(14)
        hs.font.bold = True
    else:
        hs.font.size = Pt(12)
        hs.font.bold = True


def add_cover_page():
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Project Synopsis")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Times New Roman"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(24)
    run2 = p2.add_run(
        "Sajilo Inventory: A Multitenant Inventory Management System\n"
        "Using Next.js and PostgreSQL"
    )
    run2.bold = True
    run2.font.size = Pt(16)
    run2.font.name = "Times New Roman"

    for _ in range(8):
        doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Submitted by:\nAbhay Kumar Mandal\n[College Name]\n[Date]").font.size = Pt(12)
    doc.add_page_break()


def heading(text, level=1):
    doc.add_heading(text, level=level)


def paragraph(text):
    doc.add_paragraph(text)


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def numbered(text):
    doc.add_paragraph(text, style="List Number")


FIGURE_DIR = os.path.join(os.path.dirname(__file__), "figures")

FIGURE_MAP = {
    "1": "figure_01_architecture.png",
    "2": "figure_02_dfd_l0.png",
    "3": "figure_03_dfd_l1.png",
    "4": "figure_04_er_diagram.png",
    "5": "figure_05_usecase.png",
    "6": "figure_06_class_diagram.png",
    "7": "figure_07_sequence.png",
    "8": "figure_08_gantt.png",
}


def figure_placeholder(caption):
    num = caption.split()[1].rstrip(":")
    filename = FIGURE_MAP.get(num)
    if filename:
        path = os.path.join(FIGURE_DIR, filename)
        if os.path.exists(path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(path, width=Inches(5.5))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(10)


def set_cell_shading(cell, color="D9E2F3"):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            cell.paragraphs[0].add_run(str(val)).font.size = Pt(11)
    if col_widths:
        for ri, row_obj in enumerate(table.rows):
            for ci, w in enumerate(col_widths):
                row_obj.cells[ci].width = Inches(w)
    doc.add_paragraph()
    return table


# ── BUILD DOCUMENT ─────────────────────────────────────────────────────

add_cover_page()

# ── Introduction ───────────────────────────────────────────────────────
heading("Introduction / Background")
paragraph(
    "Inventory management constitutes a fundamental operational function for "
    "retail businesses of all scales. Despite its importance, a significant "
    "proportion of small shop owners in Nepal and other developing regions "
    "continue to rely on manual record-keeping methods, including handwritten "
    "ledgers and standalone spreadsheet applications. These approaches are "
    "inherently susceptible to human error, consume considerable time, and fail "
    "to provide real-time visibility into stock levels. Consequently, businesses "
    "face recurring challenges such as overstocking, stockouts, and financial "
    "losses. As digital transformation accelerates across the small enterprise "
    "sector, there is growing demand for affordable, accessible inventory "
    "management solutions."
)
paragraph(
    "Contemporary web technologies have substantially lowered the barriers to "
    "developing sophisticated business applications that were once the exclusive "
    "domain of large enterprises with significant IT budgets. Modern frameworks "
    "enable the creation of full-stack applications with server-side rendering, "
    "secure API architectures, and type-safe database access. These advances, "
    "combined with robust relational database systems, provide the technical "
    "foundation for building scalable, multitenant platforms that can serve "
    "multiple organizations from a single deployment."
)
paragraph(
    "The proposed system, Sajilo Inventory, is a web-based, multitenant "
    "inventory management application designed to enable multiple independent "
    "shops to manage their inventory within a single platform while maintaining "
    "complete data isolation. Each tenant is provided with an independent "
    "product catalog, configurable custom attributes, stock movement tracking, "
    "and automated low-stock alerting \u2014 delivered through a modern, responsive "
    "interface tailored to the operational realities of small shop owners."
)

# ── 3. Problem Statement ──────────────────────────────────────────────
heading("Problem Statement")
paragraph(
    "Small shop owners face significant challenges managing inventory manually. "
    "Tracking product quantities, recording stock movements, monitoring expiry "
    "dates, and reconciling physical stock against records are tedious processes "
    "prone to human error. When a shop has hundreds of products, a spreadsheet-based "
    "approach becomes unmanageable. Stock discrepancies go unnoticed until a "
    "customer asks for an out-of-stock item, leading to lost sales and "
    "dissatisfied customers."
)
paragraph(
    "Existing solutions in the market are not well-suited to the needs of small "
    "shops. Enterprise resource planning (ERP) systems like Odoo and SAP are too "
    "expensive and complex. Spreadsheets lack real-time collaboration and audit "
    "trails. Off-the-shelf inventory software typically supports only a single "
    "business, forcing multi-branch owners to maintain separate systems. There "
    "is no centralized, affordable system that supports multiple shops with "
    "isolated data while providing real-time stock insights, automated alerts, "
    "and an intuitive interface."
)
p3 = doc.add_paragraph()
r3 = p3.add_run(
    "There is a clear need for an automated, multitenant, and secure inventory "
    "management system designed specifically for small shop owners."
)
r3.bold = True
r3.italic = True

# ── 4-5. Objectives ───────────────────────────────────────────────────
heading("General Objective")
paragraph(
    "To develop a secure, multitenant inventory management web application that "
    "enables small shop owners to efficiently track stock, record movements, "
    "and receive low-stock alerts."
)

heading("Specific Objectives")
objectives = [
    "To implement a multitenant architecture with complete data isolation between shops using tenant-scoped queries.",
    "To design a product catalog with custom attributes per tenant without code changes.",
    "To build a stock movement tracking system with transactional integrity and immutable audit trail.",
    "To develop low-stock and expiry alert features for proactive inventory management.",
    "To provide a stock-taking workflow for periodic reconciliation with discrepancy reporting.",
    "To deliver a responsive, mobile-friendly UI with dark mode and fast loading states.",
]
for obj in objectives:
    numbered(obj)

# ── 6. Scope ───────────────────────────────────────────────────────────
heading("Scope of the Project")
paragraph(
    "The system covers product catalog management, custom attribute definitions "
    "per tenant, stock IN/OUT recording with audit trails, low-stock and expiry "
    "alerts, stock take workflows, CSV export, and a dashboard with key metrics."
)
p = doc.add_paragraph()
r = p.add_run("Included:"); r.bold = True
for item in [
    "Google OAuth authentication with session management",
    "Tenant (shop) creation and member invitation via invite codes",
    "Product CRUD with search and pagination",
    "Custom attribute definitions (text, number, date) per tenant",
    "Stock movement recording (IN/OUT) with insufficient-stock guards",
    "Stock take lifecycle (start, count, complete, cancel)",
    "Low-stock and expiry alerts on dashboard and product detail pages",
    "CSV export for products and stock movements",
    "Dark mode, responsive design, and PWA support",
]:
    bullet(item)

p = doc.add_paragraph()
r = p.add_run("Excluded:"); r.bold = True
for item in [
    "Email/password authentication (Google OAuth only)",
    "Offline mode or local data caching",
    "Billing, invoicing, or point-of-sale (POS) integration",
    "Role-based access control beyond owner/member distinction",
    "Third-party API integrations",
    "Mobile native applications (responsive web only)",
]:
    bullet(item)

p = doc.add_paragraph()
r = p.add_run("Target Users: "); r.bold = True
p.add_run("Small shop owners and their staff in Nepal.")

# ── 7. Significance ────────────────────────────────────────────────────
heading("Significance of the Project")
for bold_part, text in [
    ("For Shop Owners: ", "Real-time stock visibility, automated alerts, and audit trails reduce errors and stockouts."),
    ("For Organizations: ", "Centralized data, better purchasing decisions, and cost-effective multitenant infrastructure."),
    ("For Society: ", "Affordable digital tools reach small businesses, supporting local economic development."),
    ("For the Developer: ", "Hands-on experience with Next.js, Prisma, PostgreSQL, OAuth, and multitenant design."),
]:
    pp = doc.add_paragraph()
    rr = pp.add_run(bold_part); rr.bold = True
    pp.add_run(text)

# ── 8. Literature Review ───────────────────────────────────────────────
heading("Literature Review")

lit_entries = [
    (
        "1. Madamidola, Daramola, Akintola, and Adeboje (2024) ",
        "conducted a comprehensive review of existing inventory management systems (IMS), "
        "tracing their evolution from manual methods through barcode scanning, RFID, and "
        "IoT-enabled solutions. Their study highlighted that while modern technologies have "
        "improved tracking accuracy, challenges remain in integrating IMS with other business "
        "processes in multi-location environments."
    ),
    (
        "2. Pedraza, Angeles, Enriquez, and Pamen (2025) ",
        "developed a cloud-based online inventory management system for SMEs with real-time "
        "stock tracking, automated low-stock alerts, role-based access control, and reporting "
        "tools. Their study demonstrated significant improvements in inventory visibility and "
        "reduction in human error compared to manual methods."
    ),
    (
        "3. Adhikari and Molla (2024) ",
        "investigated the impact of digital technology on management practices in SMEs in Nepal. "
        "Their mixed-methods study of 300 management members across seven SMEs found that "
        "digitalization enhances decision-making, customer support, and operational collaboration, "
        "while Nepali SMEs face unique barriers to digital adoption."
    ),
    (
        "4. Musuluri (2024) ",
        "explored how cloud-enabled innovation transformed inventory management and demand "
        "forecasting in retail, demonstrating that cloud integration enables real-time data "
        "processing, predictive analytics, and supply chain optimization previously limited to "
        "large enterprises."
    ),
    (
        "5. Erameh and Odoh (2021) ",
        "designed a web-based inventory control system using an SME case study, addressing "
        "overstocking, understocking, and inaccurate record-keeping at a fraction of enterprise "
        "solution costs."
    ),
    (
        "6. Balavishnu, Viswanathan, Chitradevi, Mohana Priya, and Rajkumar (2021) ",
        "developed a stock management system with automated tracking, real-time quantity "
        "monitoring, and low-stock alert mechanisms, demonstrating significant reduction in "
        "recording errors."
    ),
]
for bold_part, text in lit_entries:
    pp = doc.add_paragraph()
    rr = pp.add_run(bold_part); rr.bold = True
    pp.add_run(text)

pp = doc.add_paragraph()
rr = pp.add_run("7. Gap Identified: "); rr.bold = True
pp.add_run(
    "No existing system combines multitenancy, custom dynamic attributes, stock take workflows, "
    "automated low-stock alerts, and a modern responsive UI in a lightweight package for small "
    "Nepali shops. Sajilo Inventory fills this gap using Next.js, Prisma, and PostgreSQL."
)

# ── 9. Methodology ─────────────────────────────────────────────────────
heading("Proposed Methodology / System Approach")
paragraph(
    "The project follows the Agile development methodology with iterative two-week sprints. "
    "Agile was chosen because requirements evolved during development. Each sprint delivered "
    "a working increment: authentication, products, stock movements, dashboard, and testing."
)
pp = doc.add_paragraph()
rr = pp.add_run("Technology Stack:"); rr.bold = True
for item in [
    "Programming Language: TypeScript (primary), Python (seed scripts)",
    "Framework: Next.js 16 with App Router and React 19",
    "ORM: Prisma 7 with @prisma/adapter-pg driver adapter",
    "Database: PostgreSQL 17",
    "Authentication: Better Auth with Google OAuth provider",
    "Styling: Tailwind CSS v4 with shadcn/ui component library",
    "Validation: Zod v4 for schema validation and type inference",
    "Testing: Vitest v4 with 45 unit and isolation tests",
    "Version Control: Git",
    "Package Manager: pnpm",
    "IDE: VS Code",
]:
    bullet(item)

# ── 10. Requirements ──────────────────────────────────────────────────
heading("System Requirements")
heading("Hardware Requirements", level=2)
make_table(
    ["Component", "Minimum Specification"],
    [["Processor", "Intel Core i3 or equivalent"], ["RAM", "4 GB (8 GB recommended)"], ["Storage", "128 GB free disk space"]],
)
heading("Software Requirements", level=2)
make_table(
    ["Component", "Specification"],
    [
        ["Operating System", "Windows 10/11, Linux, or macOS"],
        ["IDE", "VS Code (or equivalent)"],
        ["Database Server", "PostgreSQL 17"],
        ["Runtime", "Node.js 20.9 or higher"],
        ["Package Manager", "pnpm 8+"],
        ["Web Browser", "Chrome, Firefox, or Edge"],
    ],
)

# ── 11. Architecture ──────────────────────────────────────────────────
heading("System Architecture / System Design")
heading("System Architecture", level=2)
paragraph(
    "Three-tier client-server architecture: React UI in the browser, Next.js server handling SSR "
    "and Server Actions, PostgreSQL accessed via Prisma ORM. A proxy.ts layer provides the first "
    "auth guard. Tenant layout enforces session validation. Every Server Action independently "
    "verifies authorization and scopes queries by tenantId."
)
figure_placeholder("Figure 1: System Architecture Diagram (3-tier with 3-layer security)")
heading("Data Flow Diagrams", level=2)
paragraph(
    "Context-level DFD (Level 0) shows Shop Owner, Shop Staff, and Google OAuth interacting with "
    "the system. Level 1 DFD decomposes into six processes: Authentication, Product Management, "
    "Stock Movement, Stock Take, Dashboard, and CSV Export."
)
figure_placeholder("Figure 2: Data Flow Diagram \u2014 Level 0 (Context)")
figure_placeholder("Figure 3: Data Flow Diagram \u2014 Level 1")
heading("Entity-Relationship Diagram", level=2)
paragraph(
    "The database consists of 11 tables: Tenant, Product, AttributeDefinition, "
    "ProductAttributeValue, StockMovement, StockTake, StockTakeItem, and Better Auth models "
    "(User, Session, Account, Verification). All tenant-scoped tables include a tenantId FK."
)
figure_placeholder("Figure 4: Entity-Relationship Diagram")
heading("UML Diagrams", level=2)
paragraph(
    "Use case diagram captures five actors and core use cases. Class diagram maps all 11 models "
    "with relationships. Sequence diagram illustrates the stock movement recording flow."
)
figure_placeholder("Figure 5: UML Use Case Diagram")
figure_placeholder("Figure 6: UML Class Diagram")
figure_placeholder("Figure 7: UML Sequence Diagram \u2014 Stock Movement")

# ── 12. Modules ────────────────────────────────────────────────────────
heading("Modules and Functional Description")
modules = [
    ("Authentication Module", "Google OAuth via Better Auth with three-layer defense-in-depth security."),
    ("Tenant Management Module", "Shop creation, invite codes, member management, owner-only operations."),
    ("Product Catalog Module", "Full CRUD with search, pagination, status badges (OK/Low/Out)."),
    ("Custom Attributes Module", "Per-tenant dynamic fields (text/number/date) with auto-rendering forms."),
    ("Stock Movement Module", "IN/OUT recording with transactional integrity and insufficient-stock guard."),
    ("Stock Take Module", "Periodic reconciliation workflow: start, count, complete with adjustments."),
    ("Dashboard Module", "Metric cards, low-stock/expiry lists, active stock take link."),
    ("CSV Export Module", "Authenticated API routes for products and movements CSV export."),
    ("Settings Module", "Attribute management, invite codes, members, financial toggle."),
]
for mod_name, mod_desc in modules:
    heading(mod_name, level=2)
    paragraph(mod_desc)

# ── 13. Deliverables ──────────────────────────────────────────────────
heading("Expected Output / Deliverables")
for d in [
    "Working web application prototype deployable on Vercel",
    "Complete TypeScript source code in a Git repository",
    "PostgreSQL database schema with migration files",
    "Seed scripts for demo data (Liquor Shop and Bicycle Shop)",
    "Project documentation suite (synopsis, architecture, database, auth, API, guides)",
    "Testing report with 45 passing tests (38 unit + 7 isolation)",
    "Screenshots of all key system screens",
    "PWA manifest and service worker for installable web experience",
]:
    bullet(d)

# ── 14. Timeline ──────────────────────────────────────────────────────
heading("Project Timeline / Gantt Chart")
paragraph("The project was developed over a 6-week period, with each phase building on the previous one.")
figure_placeholder("Figure 8: Gantt Chart \u2014 6-Week Timeline")
make_table(
    ["Phase", "Wk 1", "Wk 2", "Wk 3", "Wk 4", "Wk 5", "Wk 6"],
    [
        ["Requirement Analysis", "\u2588", "", "", "", "", ""],
        ["System Design", "", "\u2588", "", "", "", ""],
        ["Implementation", "", "", "\u2588", "\u2588", "", ""],
        ["Testing & Debugging", "", "", "", "", "\u2588", ""],
        ["Documentation & Final Submission", "", "", "", "", "", "\u2588"],
    ],
)

# ── 15. Limitations ───────────────────────────────────────────────────
heading("Limitations")
for lim in [
    "Google OAuth is the sole authentication method.",
    "Application requires active internet connection with no offline mode.",
    "Responsive web only; no native mobile app provided.",
    "Single-language interface (English only).",
    "Scalability requires additional optimization for very large data volumes.",
]:
    bullet(lim)

# ── 16. Future Enhancements ───────────────────────────────────────────
heading("Future Enhancements")
for enh in [
    "Email/password authentication as an alternative to Google OAuth.",
    "Native mobile application (React Native or Flutter).",
    "Barcode and QR code scanning for rapid product entry.",
    "AI-powered demand forecasting and auto-reorder suggestions.",
    "Multi-language support (Nepali, Hindi, etc.).",
    "Advanced analytics dashboard with interactive charts.",
    "Third-party integrations (e-commerce, accounting).",
]:
    bullet(enh)

# ── 17. Conclusion ────────────────────────────────────────────────────
heading("Conclusion")
paragraph(
    "Sajilo Inventory is a modern, multitenant inventory management web application designed "
    "to address inventory tracking challenges for small shop owners in Nepal. By combining "
    "Next.js 16, Prisma 7, PostgreSQL, and Google OAuth, the system delivers secure, real-time "
    "stock visibility with complete data isolation between tenants. The application covers product "
    "catalog management, stock movement auditing, low-stock and expiry alerts, periodic stock "
    "takes, and CSV export \u2014 all within a responsive, dark-mode-enabled interface. With 45 "
    "passing tests and a three-layer security model, Sajilo Inventory provides a robust foundation "
    "for digitizing inventory operations, with clear pathways for future growth through mobile "
    "apps, AI features, and broader language support."
)

# ── 18. References ────────────────────────────────────────────────────
heading("References")
references = [
    "Adhikari, S. N., & Molla, N. (2024). Navigating the digital shift: Exploring the impact of technology on management practices in small and medium enterprises (SMEs) in Nepal. Nepalese Journal of Management and Technology, 2(2). https://doi.org/10.3126/njmt.v2i2.68730",
    "Balavishnu, M., Viswanathan, D., Chitradevi, K., Mohana Priya, V., & Rajkumar, N. (2021). Stock management system. International Journal of Scientific Research in Computer Science, Engineering and Information Technology, 7(2), 342\u2013347. https://doi.org/10.32628/CSEIT217286",
    "Erameh, K., & Odoh, B. (2021). Design and implementation of a web-based inventory control system using a small medium enterprise (SME) as a case study. NIPES Journal of Science and Technology Research, 3(3), 211\u2013219.",
    "Madamidola, O. A., Daramola, O. A., Akintola, K., & Adeboje, O. (2024). A review of existing inventory management systems. International Journal of Research in Engineering and Science, 12(9), 40\u201350.",
    "Musuluri, C. T. (2024). Cloud-enabled innovation in retail: Transforming inventory management and demand forecasting. ResearchGate. https://doi.org/10.13140/RG.2.2.387200774",
    "Pedraza, J. S., Angeles, A. A., Enriquez, J. A., & Pamen, J. M. (2025). Online inventory management system: A web-based inventory management solution for business operations. ResearchGate. https://doi.org/10.13140/RG.2.2.391744160",
]
for ref in references:
    pp = doc.add_paragraph()
    pp.paragraph_format.left_indent = Cm(1.27)
    pp.paragraph_format.first_line_indent = Cm(-1.27)
    pp.add_run(ref).font.size = Pt(11)

doc.save(OUTPUT_PATH)
print(f"DOCX saved to: {OUTPUT_PATH}")
